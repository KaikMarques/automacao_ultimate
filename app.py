import webview
import threading
import time
import os
import sys
import platform
import subprocess
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document

# --- UTILITÁRIOS DE SISTEMA ---
def get_resource_path(relative_path):
    """Retorna o caminho absoluto para recursos (compatível com PyInstaller)."""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.getcwd(), relative_path)

# --- MÓDULO: MATRIZ DE FILTRAGEM ---
class ModuloMatriz:
    def __init__(self, api):
        self.api = api
        self.paths = {"excel": "", "word": "", "folder": ""}

    def executar(self):
        """Valida caminhos e inicia a thread."""
        if not all(self.paths.values()):
            return {"status": "error", "message": "Por favor, selecione todos os arquivos e a pasta de saída."}

        # Inicia o trabalho pesado em segundo plano para não travar a interface
        t = threading.Thread(target=self._worker)
        t.start()
        return {"status": "ok", "message": "Iniciando..."}
    
    def _worker(self):
        try:
            # 1. Trava a interface
            self.api.enviar_js('App.setLoading(true)')
            
            caminho_excel = self.paths["excel"]
            caminho_modelo = self.paths["word"]
            pasta_saida = self.paths["folder"]
            caminho_final = os.path.join(pasta_saida, "Relatorio_Consolidado_GT.docx")

            self._update_ui("Lendo Excel...", 10)
            df = pd.read_excel(caminho_excel)
            
            total = len(df)
            arquivos_temp = []

            # 2. Processa linha a linha
            for index, linha in df.iterrows():
                progresso = 10 + int((index / total) * 70)
                empresa = str(linha.get('Nome da Empresa', f'Empresa {index}'))
                self._update_ui(f"Gerando: {empresa}", progresso)

                doc = DocxTemplate(caminho_modelo)
                
                contexto = {
                    'nome_empresa': empresa,
                    'atividade':    str(linha.get('Atividade da Empresa', '')),
                    'funcionarios': str(linha.get('Funcionários', '')),
                    'gasto_anual':  self._format_moeda(linha.get('Gasto Anual', 0)),
                    'faturamento':  self._format_moeda(linha.get('Faturamento Anual', 0))
                }
                
                doc.render(contexto)
                nome_temp = os.path.join(pasta_saida, f"temp_{index}.docx")
                doc.save(nome_temp)
                arquivos_temp.append(nome_temp)

            # 3. Unificação
            if arquivos_temp:
                self._update_ui("Unificando arquivos...", 90)
                master = Document(arquivos_temp[0])
                composer = Composer(master)
                
                for arq in arquivos_temp[1:]:
                    master.add_paragraph('\n')
                    composer.append(Document(arq))
                
                composer.save(caminho_final)

                self._update_ui("Limpando temporários...", 95)
                for f in arquivos_temp:
                    try: os.remove(f)
                    except: pass

            # 4. Finalização
            self._update_ui("Finalizando...", 100)
            time.sleep(0.5)
            
            # Envia sucesso para o JS
            caminho_display = os.path.basename(caminho_final)
            self.api.enviar_js(f'App.finishSuccess("{caminho_display}")')

        except Exception as e:
            error_msg = str(e).replace('"', "'").replace('\\', '/')
            print(f"Erro: {error_msg}")
            self.api.enviar_js(f'App.showError("{error_msg}")')

    def _update_ui(self, texto, pct):
        self.api.enviar_js(f'App.updateProgress({pct}, "{texto}")')

    def _format_moeda(self, valor):
        try:
            texto = f"{float(valor):,.2f}"
            return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            return str(valor)

# --- HUB API PRINCIPAL ---
class HubApi:
    def __init__(self):
        self._window = None
        self.modulo_matriz = ModuloMatriz(self)

    def set_window(self, window):
        self._window = window

    def enviar_js(self, script):
        if self._window:
            self._window.evaluate_js(script)

    # --- Chamadas da Interface ---
    
    def selecionar_arquivo(self, modulo, tipo):
        """Abre janela de seleção de arquivo"""
        file_types = ('Excel (*.xlsx)', 'All files (*.*)') if tipo == 'excel' else ('Word (*.docx)', 'All files (*.*)')
        
        result = self._window.create_file_dialog(webview.OPEN_DIALOG, allow_multiple=False, file_types=file_types)
        
        if result:
            caminho = result[0]
            nome = os.path.basename(caminho)
            
            if modulo == 'matriz':
                self.modulo_matriz.paths[tipo] = caminho
            
            return nome
        return None

    def selecionar_pasta(self, modulo):
        """Abre janela de seleção de pasta"""
        result = self._window.create_file_dialog(webview.FOLDER_DIALOG)
        if result:
            caminho = result[0]
            nome = os.path.basename(caminho)
            
            if modulo == 'matriz':
                self.modulo_matriz.paths["folder"] = caminho
                
            return nome
        return None
    
    def abrir_pasta_saida(self):
        """Abre a pasta de saída no sistema operacional"""
        path = self.modulo_matriz.paths["folder"]
        if path and os.path.exists(path):
            if platform.system() == "Windows":
                os.startfile(path)
            elif platform.system() == "Darwin":
                subprocess.Popen(["open", path])
            else:
                subprocess.Popen(["xdg-open", path])
        else:
            self.enviar_js("alert('Caminho da pasta não encontrado.')")

    def iniciar_matriz(self):
        return self.modulo_matriz.executar()

    def fechar_app(self):
        self._window.destroy()

# --- BOOTSTRAP ---
if __name__ == '__main__':
    api = HubApi()
    
    window = webview.create_window(
        title='Grant Thornton Automation Hub',
        url='index.html',
        js_api=api,
        width=1100,
        height=720,
        resizable=True,
        background_color='#F3F4F6'
    )
    
    api.set_window(window)
    webview.start(debug=True)
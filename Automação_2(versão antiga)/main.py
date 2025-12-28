import webview
import threading
import time
import os
import sys
import pandas as pd
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document

# --- LÓGICA DE CAMINHOS ---
# Garante que funciona tanto rodando o script .py quanto o executável .exe
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.getcwd()

# --- CLASSE DE API (A Ponte entre HTML e Python) ---
class Api:
    def __init__(self):
        # CORREÇÃO 1: Usamos _window (com underline)
        # Isso esconde a janela do motor do pywebview e evita o travamento
        self._window = None
        
        self.paths = {
            "excel": "",
            "word": "",
            "folder": ""
        }

    def selecionar_arquivo(self, tipo):
        """Abre a janela nativa do Windows para selecionar arquivos"""
        file_types = ('Excel Files (*.xlsx)', 'All files (*.*)') if tipo == 'excel' else ('Word Files (*.docx)', 'All files (*.*)')
        
        # CORREÇÃO: Usar self._window
        result = self._window.create_file_dialog(
            webview.OPEN_DIALOG, 
            allow_multiple=False, 
            file_types=file_types
        )
        
        if result:
            self.paths[tipo] = result[0]
            return os.path.basename(result[0])
        return None

    def selecionar_pasta(self):
        """Abre a janela nativa para selecionar pasta"""
        # CORREÇÃO: Usar self._window
        result = self._window.create_file_dialog(webview.FOLDER_DIALOG)
        if result:
            self.paths["folder"] = result[0]
            return os.path.basename(result[0])
        return None

    def iniciar_automacao(self):
        """Chamado pelo botão 'Executar' no HTML"""
        if not all(self.paths.values()):
            return {"status": "error", "message": "Por favor, selecione todos os arquivos primeiro."}

        t = threading.Thread(target=self._worker_automacao)
        t.start()
        return {"status": "ok", "message": "Iniciando processo..."}

    def _worker_automacao(self):
        """A lógica pesada de automação"""
        try:
            # CORREÇÃO: Usar self._window
            self._window.evaluate_js('iniciarLoading()')

            caminho_excel = self.paths["excel"]
            caminho_modelo = self.paths["word"]
            pasta_saida = self.paths["folder"]
            caminho_final = os.path.join(pasta_saida, "Relatorio_Consolidado.docx")

            self._atualizar_status("Lendo Excel...", 10)
            
            # Tenta ler o Excel.
            df = pd.read_excel(caminho_excel)
            
            total_linhas = len(df)
            arquivos_temp = []

            for index, linha in df.iterrows():
                progresso = 10 + int((index / total_linhas) * 70)
                msg = f"Gerando doc {index + 1}/{total_linhas}: {linha.get('Nome da Empresa', 'Empresa')}"
                self._atualizar_status(msg, progresso)

                doc = DocxTemplate(caminho_modelo)
                
                # Mapeamento de variáveis
                contexto = {
                    'nome_empresa': linha.get('Nome da Empresa', ''),
                    'atividade':    linha.get('Atividade da Empresa', ''),
                    'funcionarios': linha.get('Funcionários', ''),
                    'gasto_anual':  self._formatar_moeda(linha.get('Gasto Anual', 0)),
                    'faturamento':  self._formatar_moeda(linha.get('Faturamento Anual', 0))
                }
                
                doc.render(contexto)
                nome_temp = os.path.join(pasta_saida, f"temp_{index}.docx")
                doc.save(nome_temp)
                arquivos_temp.append(nome_temp)

            if arquivos_temp:
                self._atualizar_status("Unificando documentos...", 90)
                master = Document(arquivos_temp[0])
                composer = Composer(master)
                
                for arq in arquivos_temp[1:]:
                    master.add_paragraph('\n')
                    composer.append(Document(arq))
                
                composer.save(caminho_final)

                self._atualizar_status("Limpando temporários...", 95)
                for f in arquivos_temp:
                    try: os.remove(f)
                    except: pass

            self._atualizar_status("Concluído!", 100)
            time.sleep(1)
            
            # Normaliza barras para evitar erro no JavaScript
            caminho_limpo = os.path.basename(caminho_final).replace('\\', '/')
            self._window.evaluate_js(f'finalizarSucesso("{caminho_limpo}")')

        except Exception as e:
            error_msg = str(e).replace('"', "'").replace('\\', '/')
            print(f"Erro Python: {error_msg}")
            # CORREÇÃO: Usar self._window
            self._window.evaluate_js(f'mostrarErro("{error_msg}")')

    def _atualizar_status(self, texto, porcentagem):
        # CORREÇÃO: Usar self._window
        self._window.evaluate_js(f'atualizarProgresso({porcentagem}, "{texto}")')

    def _formatar_moeda(self, valor):
        try:
            texto = f"{float(valor):,.2f}"
            return "R$ " + texto.replace(",", "X").replace(".", ",").replace("X", ".")
        except:
            return valor

# --- INICIALIZAÇÃO ---
if __name__ == '__main__':
    api = Api()
    
    window = webview.create_window(
        title='Grant Thornton Automation',
        url='interface.html',
        js_api=api,
        width=900,
        height=700,
        resizable=False,
        background_color='#E6E9EF'
    )
    
    # CORREÇÃO 2: Atribuir à variável com underline
    api._window = window
    
    # CORREÇÃO 3: debug=False para esconder DevTools
    webview.start(debug=False)